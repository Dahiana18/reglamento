import tkinter as tk
from tkinter import ttk
from docx import Document
from docx.oxml import OxmlElement
from docx.shared import RGBColor, Pt
from docx.oxml.ns import qn
from tkinter import messagebox

def insertar_tabla(doc, paragraph, horarios):
    # Añadir una tabla con 4 columnas y tantas filas como elementos en horarios
    table = doc.add_table(rows=1, cols=4)
    
    # Añadir encabezados
    hdr_cells = table.rows[0].cells
    headers = ['Tipo de Horario', 'Turno', 'Horario', 'Días']
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        run = hdr_cells[i].paragraphs[0].runs[0]
        run.font.name = 'Arial'
        run.font.size = Pt(11)
    
    # Añadir filas con datos
    for horario in horarios:
        row_cells = table.add_row().cells
        row_data = [horario['tipo'], horario['turno'], horario['horario'], horario['dias']]
        for i, data in enumerate(row_data):
            row_cells[i].text = data
            run = row_cells[i].paragraphs[0].runs[0]
            run.font.name = 'Arial'
            run.font.size = Pt(11)
    
    # Mover la tabla al lugar correcto
    tbl = table._tbl
    paragraph._element.addnext(tbl)

    # Añadir bordes a la tabla
    tbl = table._tbl
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')
        tblBorders.append(border)
    tbl.tblPr.append(tblBorders)
    
def capturar_orden_jerarquico(orden_jerarquico_vars):
    seleccionados = []
    for rol, var in orden_jerarquico_vars.items():
        if var.get():
            seleccionados.append(rol)
    return seleccionados

def capturar_imponer_sanciones(imponer_sanciones_vars):
    seleccionados = []
    for rol, var in imponer_sanciones_vars.items():
        if var.get():
            seleccionados.append(rol)
    return seleccionados

def reemplazar_datos_en_plantilla(nombre, municipio, departamento, objeto_social, fecha_pago, horarios, orden_jerarquico, imponer_sanciones):
    # Cargar el documento de Word
    doc = Document('template.docx')
    
    # Convertir 'nombre' a mayúsculas
    nombre = nombre.upper()

    print(f"Valor de fecha_pago: {fecha_pago}")

    # Recorrer todos los párrafos y reemplazar las palabras clave
    for p in doc.paragraphs:
        for run in p.runs:
            if "|NOMBRE|" in run.text:
                run.text = run.text.replace('|NOMBRE|', nombre)
                run.bold = True  # Aplicar negrita al run modificado

            if "|MUNICIPIO|" in run.text:
                run.text = run.text.replace('|MUNICIPIO|', municipio)
            
            if "|DEPARTAMENTO|" in run.text:
                run.text = run.text.replace('|DEPARTAMENTO|', departamento)

            if "|FECHA_PAGO|" in run.text:
                print(f"Encontrado |FECHA_PAGO| en: {run.text}")  # Mensaje de depuración
                run.text = run.text.replace('|FECHA_PAGO|', fecha_pago)
                print(f"Reemplazado por: {run.text}")  # Mensaje de depuración  

            if "|OBJETO_SOCIAL|" in run.text and isinstance(objeto_social, str):
                run.text = run.text.replace('|OBJETO_SOCIAL|', objeto_social) 
            
            if "|ORDEN_JERARQUICO|" in p.text:
                orden_jerarquico_numerado = "\n".join([f"{i+1}. {rol}" for i, rol in enumerate(orden_jerarquico)])
                p.text = p.text.replace("|ORDEN_JERARQUICO|", orden_jerarquico_numerado)
                
            if "|IMPONER_SANCIONES|" in p.text:
                imponer_sanciones_numerado = "\n".join([f"{i+1}. {rol}" for i, rol in enumerate(imponer_sanciones)])
                p.text = p.text.replace("|IMPONER_SANCIONES|", imponer_sanciones_numerado)

    for p in doc.paragraphs:
        if "|HORARIO|" in p.text:
            p.text = p.text.replace('|HORARIO|', "")
            insertar_tabla(doc, p, horarios)
            break  # Salir del bucle después de insertar la tabla

    # Guardar el documento modificado
    doc.save('documento_completado.docx')
    print("Documento generado correctamente")  # Mensaje de confirmación

# Definición de la interfaz gráfica
ventana = tk.Tk()
ventana.title("Formulario de Datos")
ventana.configure(bg='#b0d4ec')

# Crear variables de tkinter después de crear la ventana principal
operativo_var = tk.IntVar()
administrativo_var = tk.IntVar()

entry_widgets = [] 

def agregar_linea(event):
    widget = event.widget
    widget.insert(tk.END, "\n")

def agregar_fila(tipo, row=None):
    global table_frame, entry_widgets
    if row is None:
        row = len(entry_widgets) + 1
    font_settings = ("Helvetica", 14)
    entry_width = 20

    tk.Label(table_frame, text=tipo, font=font_settings).grid(row=row, column=0)
    entry_turno = tk.Entry(table_frame, font=font_settings, width=entry_width)
    entry_turno.grid(row=row, column=1)
    entry_turno.bind("<Return>", agregar_fila_manual)
    
    entry_horario = tk.Entry(table_frame, font=font_settings, width=entry_width)
    entry_horario.grid(row=row, column=2)
    entry_horario.bind("<Return>", agregar_fila_manual)
    
    entry_dias = tk.Entry(table_frame, font=font_settings, width=entry_width)
    entry_dias.grid(row=row, column=3)
    entry_dias.bind("<Return>", agregar_fila_manual)
    
    entry_widgets.append({"tipo": tipo, "entry_turno": entry_turno, "entry_horario": entry_horario, "entry_dias": entry_dias})

def generar_tabla():
    global table_frame, entry_widgets  # Asegúrate de que table_frame y entry_widgets estén accesibles
    for widget in table_frame.winfo_children():
        widget.destroy()

    entry_widgets = [] 
    
    row = 0
    font_settings = ("Helvetica", 14)
    entry_width = 20

    tk.Label(table_frame, text="Tipo de Horario", font=font_settings).grid(row=row, column=0)
    tk.Label(table_frame, text="Turno", font=font_settings).grid(row=row, column=1)
    tk.Label(table_frame, text="Horario", font=font_settings).grid(row=row, column=2)
    tk.Label(table_frame, text="Días", font=font_settings).grid(row=row, column=3)
    row += 1
    
    if operativo_var.get():
        agregar_fila("Operativo", row)
        row += 1

    if administrativo_var.get():
        agregar_fila("Administrativo", row)

# Variables globales para rastrear las últimas filas de Operativo y Administrativo
last_operativo_row = -1
last_administrativo_row = -1

def agregar_fila_manual(event=None):
    global entry_widgets, last_operativo_row, last_administrativo_row

    if operativo_var.get():
        if last_operativo_row == -1:
            agregar_fila("Operativo", 1)
            last_operativo_row = 1
        else:
            last_operativo_row += 1
            agregar_fila("Operativo", last_operativo_row + 1)
    elif administrativo_var.get():
        if last_administrativo_row == -1:
            agregar_fila("Administrativo", 1)
            last_administrativo_row = 1
        else:
            last_administrativo_row += 1
            agregar_fila("Administrativo", last_administrativo_row + 1)

def crear_formulario():  
    global nombre_entry, municipio_entry, departamento_entry, fecha_pago_entry, operativo_var, administrativo_var, objeto_social_entry, table_frame

def validar_campos():
    if not nombre_entry.get() or not municipio_entry.get() or not departamento_entry.get() or not fecha_pago_entry.get() or not objeto_social_entry.get():
        messagebox.showwarning("Campos Vacíos", "Por favor, complete todos los campos del formulario.")
        return False
    return True

def on_submit():        
    if validar_campos():
        print("Documento generado")
    nombre = nombre_entry.get()
    municipio = municipio_entry.get()
    departamento = departamento_entry.get()
    objeto_social = objeto_social_entry.get()
    fecha_pago = fecha_pago_entry.get()
    
    # Obtener el orden jerárquico
    orden_jerarquico = capturar_orden_jerarquico(orden_jerarquico_vars)
    print("Orden jerárquico seleccionado:", orden_jerarquico)
    
    # Obtener imponer sanciones
    imponer_sanciones = capturar_imponer_sanciones(imponer_sanciones_vars)
    print("imponer sanciones seleccionado:",  imponer_sanciones)


    # Construir la cadena de horarios seleccionados
    horarios = []
    for item in entry_widgets:
        horarios.append({
            "tipo": item["tipo"], 
            "turno": item["entry_turno"].get(),
            "horario": item["entry_horario"].get(),
            "dias": item["entry_dias"].get()
        })
    
    # Reemplazar datos en la plantilla
    reemplazar_datos_en_plantilla(nombre, municipio, departamento, objeto_social, fecha_pago, horarios, orden_jerarquico, imponer_sanciones)
    
    # Mostrar mensaje de confirmación
    messagebox.showinfo("Éxito", "El documento se ha generado correctamente.")

# Crear un canvas y un frame para el contenido
canvas = tk.Canvas(ventana, bg='#b0d4ec')
scroll_y = tk.Scrollbar(ventana, orient="vertical", command=canvas.yview)
scroll_x = tk.Scrollbar(ventana, orient="horizontal", command=canvas.xview)

# Frame que contendrá todos los widgets
frame_contenido = tk.Frame(canvas, bg='#b0d4ec')

# Configurar el canvas
canvas.create_window((0, 0), window=frame_contenido, anchor="nw")
canvas.configure(yscrollcommand=scroll_y.set, xscrollcommand=scroll_x.set)

# Empaquetar el canvas y la scrollbar
canvas.grid(row=0, column=0, sticky="nsew")
scroll_y.grid(row=0, column=1, sticky="ns")
scroll_x.grid(row=1, column=0, sticky="ew")

# Configurar la expansión del canvas
ventana.grid_rowconfigure(0, weight=1)
ventana.grid_columnconfigure(0, weight=1)

# Actualizar el tamaño del canvas para que se ajuste al contenido
frame_contenido.bind("<Configure>", lambda e: canvas.configure(scrollregion=canvas.bbox("all")))

font_style = ("Helvetica", 14, "italic")
bg_color = '#b0d4ec'

# Frame para los datos personales (Nombre, Municipio, Departamento)
frame_datos = tk.Frame(frame_contenido, bg=bg_color)
frame_datos.grid(padx=10, pady=10, sticky="nsew")

for i in range(6):
    frame_datos.columnconfigure(i, weight=1)

tk.Label(frame_datos, text="Nombre Empresa:", font=font_style, bg=bg_color).grid(row=0, column=0, sticky="e")
nombre_entry = tk.Entry(frame_datos, font=font_style)
nombre_entry.grid(row=0, column=1, padx=(0, 20), sticky="ew")

tk.Label(frame_datos, text="Departamento:", font=font_style, bg=bg_color).grid(row=0, column=2, sticky="e")
departamento_entry = tk.Entry(frame_datos, font=font_style)
departamento_entry.grid(row=0, column=3, padx=(0, 20), sticky="ew")

tk.Label(frame_datos, text="Municipio:", font=font_style, bg=bg_color).grid(row=0, column=4, sticky="e")
municipio_entry = tk.Entry(frame_datos, font=font_style)
municipio_entry.grid(row=0, column=5, pady=20, sticky="ew")   

tk.Label(frame_datos, text="Objeto Social:", font=font_style, bg=bg_color).grid(row=1, column=0, sticky="e")
objeto_social_entry = tk.Entry(frame_datos, font=font_style)
objeto_social_entry.grid(row=1, column=1, columnspan=5, sticky="ew", pady=20)

tk.Label(frame_datos, text="Fecha de Pago:", font=font_style, bg=bg_color).grid(row=2, column=0, sticky="e")
opciones_pago = ["los días 30 de cada mes", "los días 15 y 30 de cada mes", "catorcenales", "semanales"]
fecha_pago_entry = ttk.Combobox(frame_datos, values=opciones_pago, font=font_style, state="readonly")
fecha_pago_entry.grid(row=2, column=1, columnspan=5, sticky="ew", pady=20)
fecha_pago_entry.set("Seleccione una opción")

# Configurar la fuente del menú desplegable
ventana.option_add('*TCombobox*Listbox.font', font_style)   

# Frame para los checkbuttons (Horario de trabajo)
frame_horarios = tk.Frame(frame_contenido, bg=bg_color)
frame_horarios.grid(row=1, column=0, sticky="nsew")

# Configurar las filas para que sean responsivas
frame_horarios.rowconfigure(0, weight=1)

tk.Label(frame_horarios, text="Horario de trabajo:", bg=bg_color, font=font_style).grid(row=0, column=0, sticky="w")

operativo_cb = tk.Checkbutton(frame_horarios, text="Horario de trabajo personal operativo", variable=operativo_var, command=generar_tabla, bg=bg_color, font=font_style)
operativo_cb.grid(row=1, column=0, sticky="w")

administrativo_cb = tk.Checkbutton(frame_horarios, text="Horario de trabajo personal administrativo", variable=administrativo_var, command=generar_tabla, bg=bg_color, font=font_style)
administrativo_cb.grid(row=2, column=0, sticky="w")

# Frame para la tabla
table_frame = tk.Frame(frame_contenido)
table_frame.grid(row=3, column=0, padx=10, pady=10, sticky="nsew")

# Frame para los checkbuttons (Orden Jerárquico)
frame_orden_jerarquico = tk.Frame(frame_contenido, bg=bg_color)
frame_orden_jerarquico.grid(padx=10, pady=10, sticky="nsew")

tk.Label(frame_orden_jerarquico, text="Orden Jerárquico:", bg=bg_color, font=font_style).grid(row=0, column=0, sticky="w")

# Variables para los checkbuttons
gerente_var = tk.IntVar()
subgerente_var = tk.IntVar()
lider_talento_humano_var = tk.IntVar()
coordinador_sistemas_var = tk.IntVar()
lider_operativo_var = tk.IntVar()
supervisores_var = tk.IntVar()
operarios_manual_var = tk.IntVar()

# Diccionario para las variables de los checkbuttons
orden_jerarquico_vars = {
    "Gerente": gerente_var,
    "Subgerente": subgerente_var,
    "Líder de talento humano": lider_talento_humano_var,
    "Coordinador de sistemas integrados de gestión": coordinador_sistemas_var,
    "Líder Operativo": lider_operativo_var,
    "Supervisores": supervisores_var,
    "Operarios manual": operarios_manual_var
}

# Crear los checkbuttons
gerente_cb = tk.Checkbutton(frame_orden_jerarquico, text="Gerente", variable=gerente_var, bg=bg_color, font=font_style)
gerente_cb.grid(row=1, column=0, sticky="w")

subgerente_cb = tk.Checkbutton(frame_orden_jerarquico, text="Subgerente", variable=subgerente_var, bg=bg_color, font=font_style)
subgerente_cb.grid(row=2, column=0, sticky="w")

lider_talento_humano_cb = tk.Checkbutton(frame_orden_jerarquico, text="Líder de talento humano", variable=lider_talento_humano_var, bg=bg_color, font=font_style)
lider_talento_humano_cb.grid(row=3, column=0, sticky="w")

coordinador_sistemas_cb = tk.Checkbutton(frame_orden_jerarquico, text="Coordinador de sistemas integrados de gestión", variable=coordinador_sistemas_var, bg=bg_color, font=font_style)
coordinador_sistemas_cb.grid(row=4, column=0, sticky="w")

lider_operativo_cb = tk.Checkbutton(frame_orden_jerarquico, text="Líder Operativo", variable=lider_operativo_var, bg=bg_color, font=font_style)
lider_operativo_cb.grid(row=5, column=0, sticky="w")

supervisores_cb = tk.Checkbutton(frame_orden_jerarquico, text="Supervisores", variable=supervisores_var, bg=bg_color, font=font_style)
supervisores_cb.grid(row=6, column=0, sticky="w")

operarios_manual_cb = tk.Checkbutton(frame_orden_jerarquico, text="Operarios manual", variable=operarios_manual_var, bg=bg_color, font=font_style)
operarios_manual_cb.grid(row=7, column=0, sticky="w")

# Frame para los checkbuttons (Imponer sanciones)
frame_imponer_sanciones = tk.Frame(frame_contenido, bg=bg_color)
frame_imponer_sanciones.grid(padx=10, pady=10, sticky="nsew")

tk.Label(frame_imponer_sanciones, text="Imponer sanciones:", bg=bg_color, font=font_style).grid(row=0, column=0, sticky="w")

# Variables para los checkbuttons
gerente_var = tk.IntVar()
subgerente_var = tk.IntVar()
lider_talento_humano_var = tk.IntVar()
supervisores_var = tk.IntVar()


# Diccionario para las variables de los checkbuttons
imponer_sanciones_vars = {
    "Gerente": gerente_var,
    "Subgerente": subgerente_var,
    "Líder de talento humano": lider_talento_humano_var,    
    "Supervisores": supervisores_var,    
}
# Crear los checkbuttons
gerente_cb = tk.Checkbutton(frame_imponer_sanciones, text="Gerente", variable=gerente_var, bg=bg_color, font=font_style)
gerente_cb.grid(row=1, column=0, sticky="w")

subgerente_cb = tk.Checkbutton(frame_imponer_sanciones, text="Subgerente", variable=subgerente_var, bg=bg_color, font=font_style)
subgerente_cb.grid(row=2, column=0, sticky="w")

lider_talento_humano_cb = tk.Checkbutton(frame_imponer_sanciones, text="Líder de talento humano", variable=lider_talento_humano_var, bg=bg_color, font=font_style)
lider_talento_humano_cb.grid(row=3, column=0, sticky="w")

supervisores_cb = tk.Checkbutton(frame_imponer_sanciones, text="Supervisores", variable=supervisores_var, bg=bg_color, font=font_style)
supervisores_cb.grid(row=6, column=0, sticky="w")


# Frame para el botón de enviar
frame_botones = tk.Frame(frame_contenido, bg=bg_color)
frame_botones.grid(row=9, column=0, padx=10, pady=10, sticky="ew")

# Botón para enviar el formulario
submit_button = tk.Button(frame_botones, text="Generar Documento", command=on_submit)
submit_button.grid(row=0, column=0, pady=10, padx=10)

# Aplicar estilos al botón
submit_button.config(bg="blue", fg="white", font=("Helvetica", 12, "bold"))

# Iniciar el bucle de la aplicación Tkinter
ventana.mainloop()